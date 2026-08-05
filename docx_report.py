#!/usr/bin/env python3
"""
Gera o relatório Word de diagnóstico financeiro (python-docx puro — sem
dependência de Node/docx.js) a partir do resultado de pipeline.consolidate() +
pipeline.compute_indicators(). Os números são 100% automáticos; os "pontos de
atenção" são sinalizados por regras de referência simples (limiares usuais de
mercado) em vez de uma narrativa lapidada manualmente — para uma análise mais
fina de um caso específico (ex.: explicar uma reclassificação contábil ou uma
baixa de fornecedor), use a skill "diagnostico-bancario-grupo-cla-claudio"
com o Claude, que sabe reescrever a narrativa a cada revisão.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x78)
GRAY = RGBColor(0x59, 0x59, 0x59)
POS = RGBColor(0x2C, 0x5F, 0x2D)
CRIT = RGBColor(0x8A, 0x33, 0x24)

# Limiares de referência usados para sinalizar pontos de atenção automaticamente.
# Ajuste livremente conforme a política de crédito do seu time.
THRESHOLDS = {
    'liquidez_corrente_min': 1.0,
    'cobertura_juros_min_x': 2.0,
    'alavancagem_max_x': 3.5,
    'endividamento_geral_max_pct': 70.0,
    'concentracao_divida_cp_max_pct': 60.0,
}


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def fmt_mi(v):
    if v != v:  # NaN
        return "n/d"
    return f"R$ {v:,.1f} mi".replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_pct(v):
    if v != v:
        return "n/d"
    return f"{v:,.1f}%".replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_x(v):
    if v != v:
        return "n/d"
    return f"{v:,.2f}x".replace(",", "X").replace(".", ",").replace("X", ".")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_para(doc, text, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_indicator_table(doc, header_labels, rows, col_widths_cm=None):
    """rows: lista de listas, primeira coluna = nome do indicador (texto),
    demais colunas = valores já formatados como string."""
    table = doc.add_table(rows=1, cols=len(header_labels))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, lbl in enumerate(header_labels):
        hdr[i].text = lbl
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
        _shade_cell(hdr[i], "1F4E78")
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            for p in cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.bold = (j == 0)
            if i % 2 == 1:
                _shade_cell(cells[j], "D9E1F2")
    if col_widths_cm:
        _set_col_widths(table, col_widths_cm)
    doc.add_paragraph()
    return table


def build_report(statements, indicators, out_path, group_name="Empresa/Grupo",
                  purpose="Documento preparado para apresentação a parceiros bancários — finalidade: captação de recursos",
                  prepared_by="Consultec — Escritório de Contabilidade"):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    labels = statements['period_labels']
    n = len(labels)
    ind = indicators['por_periodo']
    agg_idx = indicators['agg_idx']
    flow_labels = labels[:-1]  # todos exceto o "Agregado"

    # ---------------- CAPA ----------------
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DIAGNÓSTICO FINANCEIRO CONSOLIDADO")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(group_name)
    r2.font.size = Pt(18)
    r2.font.bold = True
    r2.font.color.rgb = NAVY
    add_para(doc, f"Base: Balanço, DRE e DFC consolidados — {', '.join(flow_labels)}", italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, purpose, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Consolidação por soma simples das empresas do grupo, sem eliminação de saldos intercompanhias.",
              italic=True, size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ---------------- 1. SUMÁRIO EXECUTIVO ----------------
    add_heading(doc, "1. Sumário Executivo", level=1)
    crescimento = indicators.get('crescimento_pct', {})
    tese_partes = []
    if crescimento.get('receita_liquida', 0) == crescimento.get('receita_liquida', 0):
        tese_partes.append(f"a receita líquida variou {fmt_pct(crescimento['receita_liquida'])} entre o primeiro e o último período coberto")
    if crescimento.get('ebitda', 0) == crescimento.get('ebitda', 0):
        tese_partes.append(f"o EBITDA variou {fmt_pct(crescimento['ebitda'])}")
    add_para(doc,
        f"O {group_name} encerrou o período coberto por este diagnóstico com receita líquida "
        f"consolidada de {fmt_mi(ind[agg_idx]['receita_liquida_mi'])} e patrimônio líquido de "
        f"{fmt_mi(ind[agg_idx]['pl_total_mi'])}. " + ("; ".join(tese_partes) + ". " if tese_partes else "") +
        "Este relatório foi gerado automaticamente a partir da consolidação dos balancetes enviados — "
        "os números abaixo são calculados diretamente dos dados contábeis; os pontos de atenção listados "
        "na Seção 3 seguem regras de referência de mercado e devem ser revisados por quem for apresentar "
        "o documento ao banco."
    )
    add_heading(doc, "Números-chave", level=2)
    header = ["Indicador"] + labels
    rows = [
        ["Receita Líquida"] + [fmt_mi(x['receita_liquida_mi']) for x in ind],
        ["EBITDA"] + [fmt_mi(x['ebitda_mi']) for x in ind],
        ["Margem EBITDA"] + [fmt_pct(x['margem_ebitda_pct']) for x in ind],
        ["Resultado Líquido"] + [fmt_mi(x['resultado_liquido_mi']) for x in ind],
        ["Ativo Total"] + [fmt_mi(x['ativo_total_mi']) for x in ind],
        ["Patrimônio Líquido"] + [fmt_mi(x['pl_total_mi']) for x in ind],
    ]
    add_indicator_table(doc, header, rows)

    # ---------------- 2. INDICADORES DETALHADOS ----------------
    add_heading(doc, "2. Indicadores Financeiros Detalhados", level=1)

    add_heading(doc, "2.1. Liquidez e margens", level=2)
    add_indicator_table(doc, header, [
        ["Liquidez Corrente"] + [fmt_x(x['liquidez_corrente']) for x in ind],
        ["Margem Bruta"] + [fmt_pct(x['margem_bruta_pct']) for x in ind],
        ["Margem EBITDA"] + [fmt_pct(x['margem_ebitda_pct']) for x in ind],
        ["Margem Líquida"] + [fmt_pct(x['margem_liquida_pct']) for x in ind],
    ])

    add_heading(doc, "2.2. Endividamento e alavancagem", level=2)
    add_indicator_table(doc, header, [
        ["Endividamento Geral (Passivo/Ativo)"] + [fmt_pct(x['endividamento_geral_pct']) for x in ind],
        ["Patrimônio Líquido / Ativo Total"] + [fmt_pct(x['pl_sobre_ativo_pct']) for x in ind],
        ["Dívida Financeira Bruta"] + [fmt_mi(x['divida_bruta_mi']) for x in ind],
        ["Dívida Financeira Líquida"] + [fmt_mi(x['divida_liquida_mi']) for x in ind],
        ["Concentração de dívida no curto prazo"] + [fmt_pct(x['concentracao_divida_cp_pct']) for x in ind],
        ["Cobertura de Juros (EBITDA/Desp. Financeiras)"] + [fmt_x(x['cobertura_juros_x']) for x in ind],
    ])
    add_para(doc,
        f"Alavancagem (Dívida Líquida / EBITDA anualizado, fator de anualização "
        f"{indicators['fator_anualizacao']:.2f}x sobre o período coberto): "
        f"{fmt_x(indicators['alavancagem_x'])}. ROE anualizado: {fmt_pct(indicators['roe_anualizado_pct'])}. "
        f"ROA anualizado: {fmt_pct(indicators['roa_anualizado_pct'])}.",
        italic=True, size=10, color=GRAY,
    )

    add_heading(doc, "2.3. Fluxo de caixa (DFC)", level=2)
    add_indicator_table(doc, header, [
        ["Caixa Operacional"] + [fmt_mi(x['caixa_operacional_mi']) for x in ind],
        ["Caixa de Investimento"] + [fmt_mi(x['caixa_investimento_mi']) for x in ind],
        ["Caixa de Financiamento"] + [fmt_mi(x['caixa_financiamento_mi']) for x in ind],
        ["Variação de Caixa no Período"] + [fmt_mi(x['variacao_caixa_mi']) for x in ind],
    ])

    # ---------------- 3. PONTOS DE ATENÇÃO (regras automáticas) ----------------
    add_heading(doc, "3. Pontos de Atenção (sinalização automática)", level=1)
    add_para(doc,
        "Os itens abaixo foram sinalizados comparando o último período coberto com os limiares de "
        "referência configurados neste relatório — eles NÃO substituem uma análise qualitativa (ex.: "
        "explicar uma reclassificação contábil específica ou uma baixa de fornecedor). Ajuste os "
        "limiares em THRESHOLDS conforme a política de crédito do seu time."
    )
    last = ind[indicators['last_flow_idx']]
    alertas = []
    if last['liquidez_corrente'] < THRESHOLDS['liquidez_corrente_min']:
        alertas.append(("Liquidez Corrente abaixo de 1,0x",
                         f"Liquidez Corrente de {fmt_x(last['liquidez_corrente'])} no último período — "
                         f"o ativo circulante não cobre integralmente o passivo circulante."))
    if last['cobertura_juros_x'] < THRESHOLDS['cobertura_juros_min_x']:
        alertas.append(("Cobertura de juros apertada",
                         f"EBITDA cobre apenas {fmt_x(last['cobertura_juros_x'])} das despesas financeiras "
                         f"do último período — referência de conforto usual dos bancos é acima de "
                         f"{THRESHOLDS['cobertura_juros_min_x']:.1f}x."))
    if indicators['alavancagem_x'] == indicators['alavancagem_x'] and indicators['alavancagem_x'] > THRESHOLDS['alavancagem_max_x']:
        alertas.append(("Alavancagem acima da referência",
                         f"Dívida Líquida / EBITDA anualizado em {fmt_x(indicators['alavancagem_x'])} — "
                         f"referência de conforto usual é até {THRESHOLDS['alavancagem_max_x']:.1f}x."))
    if last['endividamento_geral_pct'] > THRESHOLDS['endividamento_geral_max_pct']:
        alertas.append(("Endividamento geral elevado",
                         f"Passivo total representa {fmt_pct(last['endividamento_geral_pct'])} do Ativo Total "
                         f"no último período."))
    if last['concentracao_divida_cp_pct'] > THRESHOLDS['concentracao_divida_cp_max_pct']:
        alertas.append(("Concentração de dívida no curto prazo",
                         f"{fmt_pct(last['concentracao_divida_cp_pct'])} do passivo total vence no curto prazo — "
                         f"pressiona o caixa e aumenta a dependência de rolagem de dívida."))
    if last['caixa_operacional_mi'] < 0:
        alertas.append(("Caixa operacional negativo no último período",
                         f"A operação consumiu {fmt_mi(abs(last['caixa_operacional_mi']))} de caixa no último "
                         f"período coberto — investigue se é crescimento de capital de giro, baixa de "
                         f"fornecedor deliberada, ou sinal de aperto financeiro antes de apresentar ao banco."))
    if not alertas:
        add_para(doc, "Nenhum indicador ultrapassou os limiares de referência configurados no último período coberto.")
    else:
        for titulo, texto in alertas:
            h = doc.add_paragraph()
            r = h.add_run(f"⚠ {titulo}: ")
            r.font.bold = True
            r.font.color.rgb = CRIT
            r.font.size = Pt(11)
            r2 = h.add_run(texto)
            r2.font.size = Pt(11)

    # ---------------- 4. NOTA METODOLÓGICA ----------------
    add_heading(doc, "4. Nota Metodológica", level=1)
    add_para(doc,
        "Os números deste diagnóstico vêm da consolidação por soma simples (sem eliminação de saldos "
        "intercompanhias) das empresas do grupo, com base nos balancetes de verificação enviados. "
        "Indicadores que dependem de posição (Balanço) usam o saldo de fim de período; indicadores de "
        "fluxo (DRE, DFC) somam os períodos quando aplicável. Indicadores anualizados (ROE, ROA, Dívida "
        f"Líquida/EBITDA) foram obtidos multiplicando o resultado do período coberto por um fator de "
        f"anualização ({indicators['fator_anualizacao']:.2f}x) — uma aproximação simples que não considera "
        "sazonalidade, e deve ser tratada como estimativa. Relatório gerado automaticamente; recomenda-se "
        "revisão humana antes do envio ao banco, especialmente para explicar qualquer reclassificação "
        "contábil, baixa de fornecedor ou correção pontual de saldo ocorrida no período."
    )
    add_para(doc, f"Elaborado com apoio de {prepared_by}.", italic=True, size=9, color=GRAY)

    doc.save(out_path)
    return out_path
